"""Export helpers for Markdown, PDF, and DOCX downloads."""

from __future__ import annotations

import html
import re
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import BaseDocTemplate, Frame, PageTemplate, Paragraph, Spacer, Table, TableStyle


def export_to_markdown(markdown_text: str) -> bytes:
    """Return the markdown file payload."""
    return markdown_text.encode("utf-8")


def export_to_pdf(markdown_text: str, title: str = "Cheat Sheet") -> bytes:
    """Render a compact, two-column PDF from markdown-like content."""
    buffer = BytesIO()
    doc = BaseDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=10 * mm,
        rightMargin=10 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
        title=title,
    )

    column_gap = 6 * mm
    usable_width = A4[0] - doc.leftMargin - doc.rightMargin
    frame_width = (usable_width - column_gap) / 2
    frame_height = A4[1] - doc.topMargin - doc.bottomMargin

    frames = [
        Frame(doc.leftMargin, doc.bottomMargin, frame_width, frame_height, id="left"),
        Frame(doc.leftMargin + frame_width + column_gap, doc.bottomMargin, frame_width, frame_height, id="right"),
    ]
    doc.addPageTemplates([PageTemplate(id="TwoColumn", frames=frames)])

    font_name = _resolve_pdf_font(markdown_text)
    styles = _build_pdf_styles(font_name)
    story = _markdown_to_pdf_flowables(markdown_text, styles, frame_width)
    doc.build(story)

    return buffer.getvalue()


def export_to_docx(markdown_text: str, title: str = "Cheat Sheet") -> bytes:
    """Render a Word document while keeping headings, bullets, and tables readable."""
    document = Document()
    document.core_properties.title = title
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(9)

    heading_1 = document.styles["Heading 1"]
    heading_1.font.name = "Arial"
    heading_1.font.size = Pt(13)

    heading_2 = document.styles["Heading 2"]
    heading_2.font.name = "Arial"
    heading_2.font.size = Pt(11)

    heading_3 = document.styles["Heading 3"]
    heading_3.font.name = "Arial"
    heading_3.font.size = Pt(10)

    blocks = _parse_markdown_blocks(markdown_text)

    for block in blocks:
        if block["type"] == "heading":
            level = min(block["level"], 3)
            document.add_heading(block["text"], level=level)
        elif block["type"] == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.add_run(block["text"])
        elif block["type"] == "paragraph":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run(block["text"])
        elif block["type"] == "table":
            rows = block["rows"]
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for row_index, row in enumerate(rows):
                for col_index, cell_text in enumerate(row):
                    table.cell(row_index, col_index).text = cell_text
            document.add_paragraph("")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _resolve_pdf_font(markdown_text: str) -> str:
    if any(ord(char) > 127 for char in markdown_text):
        try:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            return "STSong-Light"
        except Exception:
            return "Helvetica"
    return "Helvetica"


def _build_pdf_styles(font_name: str) -> dict[str, ParagraphStyle]:
    stylesheet = getSampleStyleSheet()
    base_leading = 8.5 if font_name == "Helvetica" else 9.2

    return {
        "title": ParagraphStyle(
            "TitleCompact",
            parent=stylesheet["Title"],
            fontName=font_name,
            fontSize=12,
            leading=13,
            spaceAfter=4,
            textColor=colors.black,
        ),
        "h2": ParagraphStyle(
            "Heading2Compact",
            parent=stylesheet["Heading2"],
            fontName=font_name,
            fontSize=9.2,
            leading=10.5,
            spaceBefore=4,
            spaceAfter=2,
            textColor=colors.black,
        ),
        "h3": ParagraphStyle(
            "Heading3Compact",
            parent=stylesheet["Heading3"],
            fontName=font_name,
            fontSize=8.4,
            leading=9.2,
            spaceBefore=2,
            spaceAfter=1,
            textColor=colors.black,
        ),
        "body": ParagraphStyle(
            "BodyCompact",
            parent=stylesheet["BodyText"],
            fontName=font_name,
            fontSize=7.2,
            leading=base_leading,
            spaceAfter=2,
            textColor=colors.black,
        ),
        "bullet": ParagraphStyle(
            "BulletCompact",
            parent=stylesheet["BodyText"],
            fontName=font_name,
            fontSize=7.2,
            leading=base_leading,
            leftIndent=8,
            firstLineIndent=0,
            bulletIndent=0,
            spaceAfter=1,
            textColor=colors.black,
        ),
    }


def _markdown_to_pdf_flowables(markdown_text: str, styles: dict[str, ParagraphStyle], frame_width: float) -> list:
    flowables: list = []

    for block in _parse_markdown_blocks(markdown_text):
        if block["type"] == "heading":
            if block["level"] == 1:
                flowables.append(Paragraph(_inline_to_reportlab_html(block["text"]), styles["title"]))
            elif block["level"] == 2:
                flowables.append(Paragraph(_inline_to_reportlab_html(block["text"]), styles["h2"]))
            else:
                flowables.append(Paragraph(_inline_to_reportlab_html(block["text"]), styles["h3"]))
            continue

        if block["type"] == "bullet":
            flowables.append(
                Paragraph(
                    _inline_to_reportlab_html(block["text"]),
                    styles["bullet"],
                    bulletText="•",
                )
            )
            continue

        if block["type"] == "paragraph":
            flowables.append(Paragraph(_inline_to_reportlab_html(block["text"]), styles["body"]))
            continue

        if block["type"] == "table":
            rows = [
                [Paragraph(_inline_to_reportlab_html(cell), styles["body"]) for cell in row]
                for row in block["rows"]
            ]
            column_width = frame_width / max(1, len(rows[0]))
            table = Table(rows, colWidths=[column_width] * len(rows[0]), repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EFEFEF")),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#666666")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 3),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                        ("TOPPADDING", (0, 0), (-1, -1), 2),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                    ]
                )
            )
            flowables.append(table)
            flowables.append(Spacer(1, 2))

    return flowables


def _parse_markdown_blocks(markdown_text: str) -> list[dict]:
    lines = markdown_text.splitlines()
    blocks: list[dict] = []
    index = 0

    while index < len(lines):
        raw_line = lines[index].rstrip()
        stripped = raw_line.strip()

        if not stripped:
            index += 1
            continue

        if stripped.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = _parse_markdown_table(table_lines)
            if rows:
                blocks.append({"type": "table", "rows": rows})
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            blocks.append(
                {
                    "type": "heading",
                    "level": len(heading_match.group(1)),
                    "text": heading_match.group(2).strip(),
                }
            )
            index += 1
            continue

        bullet_match = re.match(r"^([-*]|\d+\.)\s+(.*)$", stripped)
        if bullet_match:
            blocks.append({"type": "bullet", "text": bullet_match.group(2).strip()})
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            lookahead = lines[index].strip()
            if not lookahead:
                break
            if lookahead.startswith("|") or re.match(r"^(#{1,6})\s+", lookahead) or re.match(r"^([-*]|\d+\.)\s+", lookahead):
                break
            paragraph_lines.append(lookahead)
            index += 1

        blocks.append({"type": "paragraph", "text": " ".join(paragraph_lines)})

    return blocks


def _parse_markdown_table(table_lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if not cells:
            continue
        if all(re.fullmatch(r"[:\-\s]+", cell) for cell in cells):
            continue
        rows.append(cells)

    if not rows:
        return rows

    width = max(len(row) for row in rows)
    return [row + [""] * (width - len(row)) for row in rows]


def _inline_to_reportlab_html(text: str) -> str:
    escaped = html.escape(text)
    escaped = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(r"`(.+?)`", r'<font face="Courier">\1</font>', escaped)
    return escaped
